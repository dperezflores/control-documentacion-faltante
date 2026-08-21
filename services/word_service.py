from __future__ import annotations

from collections import defaultdict
from io import BytesIO

from docx import Document


def build_request_docx(requirement: str, cuts: list[int], records: list[dict]) -> bytes:
    doc = Document()
    doc.add_heading("Solicitud de documentación", level=1)
    doc.add_paragraph(f"Requerimiento de referencia: {requirement}")
    doc.add_paragraph("Cortes incluidos: " + ", ".join(str(c) for c in sorted(cuts)))
    doc.add_paragraph("Se solicita proporcionar la siguiente documentación:")

    grouped = defaultdict(list)
    for item in records:
        if item["documento"] not in grouped[item["contrato"]]:
            grouped[item["contrato"]].append(item["documento"])

    for contract, docs in grouped.items():
        doc.add_heading(contract, level=2)
        for text in docs:
            doc.add_paragraph(text, style="List Number")

    doc.add_paragraph("\nNota: este documento corresponde al prototipo. La plantilla institucional se incorporará posteriormente.")
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
